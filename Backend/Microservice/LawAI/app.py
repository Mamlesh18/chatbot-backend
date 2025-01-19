from flask import Flask, jsonify, request, send_file
from prometheus_flask_exporter import PrometheusMetrics
import google.generativeai as genai
from docx import Document
import requests
import io
from flask_cors import CORS
from prometheus_client import generate_latest, CONTENT_TYPE_LATEST



app = Flask(__name__)
# Initialize Prometheus metrics
metrics = PrometheusMetrics(app, defaults_prefix='my_app')

# Enable CORS for all routes
CORS(app)


@app.route('/metrics')
def metrics_route():

    return generate_latest(), 200, {'Content-Type': CONTENT_TYPE_LATEST}

class GeminiAI:
    def __init__(self, api_key, model_name):
        self.api_key = api_key
        self.model_name = model_name
        genai.configure(api_key=self.api_key)

    def generate_response(self, prompt):
        try:
            model = genai.GenerativeModel(self.model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"Error occurred: {e}"

@app.route('/searchgeminilaw', methods=['POST'])
def geminirest():
    print("1 ------------ question")

    data = request.get_json()
    query = data.get('query', '')
    print("3 ------------ question")

    chat_prompt = (
     """
        RESIDENTIAL RENTAL AGREEMENT
        This agreement made at #city, #state on this #ddmmyy between #landlordname, residing at #landlordaddress1, #lordaddressline2, #lordcity, #lordstate, #lordpincode hereinafter referred to as the `LESSOR` of the One Part AND #tenantname, residing at  #tenantaddress1, #tenantaddressline2, #tencity, #tenstate, #tenpincode hereinafter referred to as the `LESSEE` of the other Part;
        WHEREAS the Lessor is the lawful owner of, and otherwise well sufficiently entitled to #leasepropertyaddress1, #leaseaddressline2, #leasecity, #leasestate, #leasepincode falling in the category, #independenthouse / #apartment / #farmhouse / #residentialproperty and comprising of #xbedrooms, #xbathrooms, #xcarparks with an extent of #xxxxsquarefeet hereinafter referred to as the `said premises`. 
        AND WHEREAS at the request of the Lessee, the Lessor has agreed to let the said premises to the tenant for a term of #leaseterm commencing from #leasestartdate in the manner hereinafter appearing. 

        NOW THIS AGREEMENT WITNESSETH AND IT IS HEREBY AGREED BY AND BETWEEN THE PARTIES AS UNDER:
        1.	That the Lessor hereby grants to the Lessee, the right to enter into use and remain in the said premises along with the existing fixtures and fittings listed in Annexure 1 to this Agreement and that the Lessee shall be entitled to peacefully possess, and enjoy possession of the said premises, and the other rights herein.
        2.	That the lease hereby granted shall, unless cancelled earlier under any provision of this Agreement, remain in force for a period of #leaseterm. 
        3.	That the Lessee will have the option to terminate this lease by giving #onemonthnotice in writing to the Lessor.
        4.	That the Lessee shall have no right to create any sub-lease or assign or transfer in any manner the lease or give to anyone the possession of the said premises or any part thereof.
        5.	That the Lessee shall use the said premises only for residential purposes.
        6.	That the Lessor shall, before handing over the said premises, ensure the working of sanitary, electrical and water supply connections and other fittings pertaining to the said premises. It is agreed that it shall be the responsibility of the Lessor for their return in the working condition at the time of re-possession of the said premises (reasonable wear and tear and loss or damage by fire, flood, rains, accident, irresistible force or act of God excepted).
        7.	That the Lessee is not authorized to make any alteration in the construction of the said premises. The Lessee may however install and remove his own fittings and fixtures, provided this is done without causing any excessive damage or loss to the said premises.
        8.	That the day-to-day repair jobs such as fuse blow out, replacement of light bulbs/tubes, leakage of water taps, maintenance of the water pump and other minor repairs, etc., shall be effected by the Lessee at its own cost, and any major repairs, either structural or to the electrical or water connection, plumbing leaks, water seepage shall be attended to by the Lessor. In the event of the Lessor failing to carry out the repairs on receiving notice from the Lessee, the Lessee shall undertake the necessary repairs and the Lessor will be liable to immediately reimburse costs incurred by the Lessee.
        9.	That the Lessor or its duly authorized agent shall have the right to enter into or upon the said premises or any part thereof at a mutually arranged convenient time for the purpose of inspection. 
        10.	That the Lessee shall use the said premises along with its fixtures and fitting in careful and responsible manner and shall handover the premises to the Lessor in working condition (reasonable wear and tear and loss or damage by fire, flood, rains, accidents, irresistible force or act of God excepted).
        11.	That in consideration of use of the said premises the Lessee agrees that he shall pay to the Lessor during the period of this agreement, a monthly rent at the rate of #monthlyrentalinnumber&words. The amount will be paid in advance on or before the date of #paiday of every English calendar month.
        12.	It is hereby agreed that if default is made by the lessee in payment of the rent for a period of three months, or in observance and performance of any of the covenants and stipulations hereby contained and on the part to be observed and performed by the lessee, then on such default, the lessor shall be entitled in addition to or in the alternative to any other remedy that may be available to him at this discretion, to terminate the lease and eject the lessee from the said premises; and to take possession thereof as full and absolute owner thereof, provided that a notice in writing shall be given by the lessor to the lessee of his intention to terminate the lease and to take possession of the said premises. If the arrears of rent are paid or the lessee comply with or carry out the covenants and conditions or stipulations, within fifteen days from the service of such notice, then the lessor shall not be entitled to take possession of the said premises.
        13.	That in addition to the compensation mentioned above, the Lessee shall pay the actual electricity, shared maintenance, water bills for the period of the agreement directly to the authorities concerned. The relevant `start date` meter readings are #startingmetereading. 
        14.	That the Lessee has paid to the Lessor a sum of #rentaldepositinumber&words as deposit, free of interest, which the Lessor does accept and acknowledge. This deposit is for the due performance and observance of the terms and conditions of this Agreement. The deposit shall be returned to the Lessee simultaneously with the Lessee vacating the said premises. In the event of failure on the part of the Lessor to refund the said deposit amount to the Lessee as aforesaid, the Lessee shall be entitled to continue to use and occupy the said premises without payment of any rent until the Lessor refunds the said amount (without prejudice to the Lessee`s rights and remedies in law to recover the deposit).
        15.	That the Lessor shall be responsible for the payment of all taxes and levies pertaining to the said premises including but not limited to House Tax, Property Tax, other cesses, if any, and any other statutory taxes, levied by the Government or Governmental Departments. During the term of this Agreement, the Lessor shall comply with all rules, regulations and requirements of any statutory authority, local, state and central government and governmental departments in relation to the said premises.
        IN WITNESS WHEREOF, the parties hereto have set their hands on the day and year first hereinabove mentioned. 

        Lessor,	Lessee,
        #name	#name
        # landlordaddress1	# tenantaddress1
        #lordaddressline2	#tenantaddressline2
        #lordcity, #lordstate, #lordpincode	#tencity, #tenstate, #tenpincode


        WITNESS ONE	WITNESS TWO


        [Name & Address]	[Name & Address]

        ANNEXURE I
        List of fixtures and fittings provided in #leasepropertyaddress1, #leaseaddressline2, #leasecity, #leasestate, #leasepincode: 
        1.	#item1
        2.	#item2
        3.	#item3
        """
        f"You are an assistant tasked with filling out a Residential Rental Agreement based on the user's query. The agreement contains placeholders that need to be replaced with specific details from the user's input. The placeholders are as in #city like this, itll start with a # so you need to replace those fileds only and return the full aggremmentYour task is to process the following user query and extract the required information to replace the placeholders in the agreement accordingly. If any field s missing raise a question to user about please provide that specific field The query is: {query}"
)

    print("6 ------------ question")

    api_key = "AIzaSyDcP3_6sDB3P8lZkIyv0YSeFfvMsh_5RsQ"
    model_name = 'gemini-1.5-flash-latest'
    gemini_client = GeminiAI(api_key, model_name)
    print("7 ------------ question")

    response = gemini_client.generate_response(chat_prompt)
    print("8 ------------ question")
    if isinstance(response, str):

        points = response.split('\n')  # Split by new lines or you can use regex for better splitting
        # Clean up each point (remove extra spaces)
        points = [point.strip() for point in points if point.strip()]

    return jsonify({"answer": points})
def update_document(city,state,date,landlordname,landlordaddress1,lordaddressline2,lordcity,lordstate,lordpincode,
        tenantname,tenantaddress1,tenantaddressline2,tencity,tenstate,tenpincode,leasepropertyaddress1,
        leaseaddressline2,leasecity,leasestate,leasepincode,category,xbedrooms,xbathrooms,xcarparks,
        xxxxsquarefeet,leaseterm,leasestartdate,onemonthnotice,monthlyrentalintextwords,startingmetereading,
        rentaldeposititextwords,item1,item2,item3):
    # Download the document from the URL
    document_url = "https://docs.google.com/document/d/10uUedusfVfjgl15LNeXSn6Cosh7-4S8d/export?format=docx"
   
    response = requests.get(document_url)
    # Check if the request was successful
    if response.status_code == 200:
        # Load the downloaded document
        doc = Document(io.BytesIO(response.content))

        # Replace placeholders with values
        for paragraph in doc.paragraphs:
            if '#city' in paragraph.text:
                paragraph.text = paragraph.text.replace('#city', city)
            if '#state' in paragraph.text:
                paragraph.text = paragraph.text.replace('#state', state)
            if '#ddmmyy' in paragraph.text:
                paragraph.text = paragraph.text.replace('#ddmmyy', date)
            if '#landlordname' in paragraph.text:
                paragraph.text = paragraph.text.replace('#landlordname', landlordname)
            if '#landlordaddress1' in paragraph.text:
                paragraph.text = paragraph.text.replace('#landlordaddress1', landlordaddress1)
            if '#lordaddressline2' in paragraph.text:
                paragraph.text = paragraph.text.replace('#lordaddressline2', lordaddressline2)
            if '#lordcity' in paragraph.text:
                paragraph.text = paragraph.text.replace('#lordcity', lordcity)
            if '#lordstate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#lordstate', lordstate)
            if '#lordpincode ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#lordpincode', lordpincode)
            if '#tenantname' in paragraph.text:
                paragraph.text = paragraph.text.replace('#tenantname', tenantname)
            if '#tenantaddress1' in paragraph.text:
                paragraph.text = paragraph.text.replace('#tenantaddress1', tenantaddress1)
            if '#tenantaddressline2' in paragraph.text:
                paragraph.text = paragraph.text.replace('#tenantaddressline2', tenantaddressline2)
            if '#tencity' in paragraph.text:
                paragraph.text = paragraph.text.replace('#tencity', tencity)
            if '#tenstate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#tenstate', tenstate)
            if '#tenpincode ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#tenpincode ', tenpincode)
            if '#leasepropertyaddress1' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leasepropertyaddress1', leasepropertyaddress1)
            if '#leaseaddressline2' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leaseaddressline2', leaseaddressline2)
            if '#leasecity' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leasecity', leasecity)
            if '#leasestate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leasestate', leasestate)
            if '#leasepincode ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leasepincode', leasepincode)
            if '#category ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#category', category)
            if '#xbedrooms' in paragraph.text:
                paragraph.text = paragraph.text.replace('#xbedrooms', xbedrooms)
            if '#xbathrooms' in paragraph.text:
                paragraph.text = paragraph.text.replace('#xbathrooms', xbathrooms)
            if '#xcarparks' in paragraph.text:
                paragraph.text = paragraph.text.replace('#xcarparks', xcarparks)
            if '#xxxxsquarefeet' in paragraph.text:
                paragraph.text = paragraph.text.replace('#xxxxsquarefeet', xxxxsquarefeet)
            if '#leaseterm' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leaseterm', leaseterm)
            if '#leasestartdate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#leasestartdate', leasestartdate)
            if '#onemonthnotice' in paragraph.text:
                paragraph.text = paragraph.text.replace('#onemonthnotice', onemonthnotice)
            if '#monthlyrentalinnumber&words' in paragraph.text:
                paragraph.text = paragraph.text.replace('#monthlyrentalinnumber&words', monthlyrentalintextwords)
            if '#startingmetereading' in paragraph.text:
                paragraph.text = paragraph.text.replace('#startingmetereading', startingmetereading)
            if '#rentaldepositinumber&words' in paragraph.text:
                paragraph.text = paragraph.text.replace('#rentaldepositinumber&words', rentaldeposititextwords)
            if '#item1' in paragraph.text:
                paragraph.text = paragraph.text.replace('#item1', item1)
            if '#item2' in paragraph.text:
                paragraph.text = paragraph.text.replace('#item2', item2)
            if '#item3' in paragraph.text:
                paragraph.text = paragraph.text.replace('#item3', item3)
            
        
        return doc
    else:
        return None

@app.route('/print_info', methods=['POST'])
def print_info():
    data = request.json
    city = data.get('city')
    state = data.get('state')
    date = data.get('date')
    landlordname = data.get('landlordname')
    landlordaddress1 = data.get('landlordaddress1')
    lordaddressline2 = data.get('lordaddressline2')
    lordcity = data.get('lordcity')
    lordstate = data.get('lordstate')
    lordpincode = data.get('lordpincode')
    tenantname = data.get('tenantname')
    tenantaddress1 = data.get('tenantaddress1')
    tenantaddressline2 = data.get('tenantaddressline2')
    tencity = data.get('tencity')
    tenstate = data.get('tenstate')
    tenpincode = data.get('tenpincode')
    leasepropertyaddress1 = data.get('leasepropertyaddress1')
    leaseaddressline2 = data.get('leaseaddressline2')
    leasecity = data.get('leasecity')
    leasestate = data.get('leasestate')
    leasepincode = data.get('leasepincode')
    category = data.get('category')
    xbedrooms = data.get('xbedrooms')
    xbathrooms = data.get('xbathrooms')
    xcarparks = data.get('xcarparks')
    xxxxsquarefeet = data.get('xxxxsquarefeet')
    leaseterm = data.get('leaseterm')
    leasestartdate = data.get('leasestartdate')
    onemonthnotice = data.get('onemonthnotice')
    monthlyrentalintextwords = data.get('monthlyrentalintextwords')
    startingmetereading = data.get('startingmetereading')
    rentaldeposititextwords = data.get('rentaldeposititextwords')
    item1 = data.get('item1')
    item2 = data.get('item2')
    item3 = data.get('item3')
  
    
    
    # Update the document with the received values
    updated_doc = update_document(city,state,date,landlordname,landlordaddress1,lordaddressline2,lordcity,lordstate,lordpincode,
        tenantname,tenantaddress1,tenantaddressline2,tencity,tenstate,tenpincode,leasepropertyaddress1,
        leaseaddressline2,leasecity,leasestate,leasepincode,category,xbedrooms,xbathrooms,xcarparks,
        xxxxsquarefeet,leaseterm,leasestartdate,onemonthnotice,monthlyrentalintextwords,startingmetereading,
        rentaldeposititextwords,item1,item2,item3)
    
    if updated_doc:
        # Send the updated document back to the React app
        updated_doc.save('updated_document.docx')
        return send_file('updated_document.docx', as_attachment=True)
    else:
        return "Failed to fetch the document from the URL", 500
    


def update_document_Divorce(pet1name,pet1age,pet1occupation,pet1address,pet1mobileNo,pet1emailid,pet2name,pet2age,
                    pet2occupation,pet2address,pet2mobileNo,pet2emailid,marriedplace,marrieddate,religion,registarplace,
                    pet1premarstatus,pet2premarstatus,noofchildren,childname,childage,childdob,childcustody,
                    section,act,caseno,idproof,marriageproof,residentialproof,propertydocument,receipt ):


    # Download the document from the URL
    document_url = "https://docs.google.com/document/d/1VfzQThWR0Kfw-jqd-ys_wzb_OlzKENXU/export?format=docx"
    response = requests.get(document_url)
    
    # Check if the request was successful
    if response.status_code == 200:
        # Load the downloaded document
        doc = Document(io.BytesIO(response.content))

        # Replace placeholders with values
        for paragraph in doc.paragraphs:
            if '#pet1name' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1name', pet1name)
            if '#pet1age' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1age', pet1age)
            if '#pet1occupation' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1occupation', pet1occupation)
            if '#pet1address' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1address', pet1address)
            if '#pet1mobileNo' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1mobileNo', pet1mobileNo)
            if '#pet1emailid' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1emailid', pet1emailid)
            if '#pet2name' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2name', pet2name)
            if '#pet2age' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2age', pet2age)
            if '#pet2occupation ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2occupation', pet2occupation)
            if '#pet2address' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2address', pet2address)
            if '#pet2mobileNo' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2mobileNo', pet2mobileNo)
            if '#pet2emailid' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2emailid', pet2emailid)
            if '#marriedplace' in paragraph.text:
                paragraph.text = paragraph.text.replace('#marriedplace', marriedplace)
            if '#marrieddate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#marrieddate', marrieddate)
            if '#religion ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#religion ', religion)
            if '#registarplace' in paragraph.text:
                paragraph.text = paragraph.text.replace('#registarplace', registarplace)
            if '#pet1premarstatus' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet1premarstatus', pet1premarstatus)
            if '#pet2premarstatus' in paragraph.text:
                paragraph.text = paragraph.text.replace('#pet2premarstatus', pet2premarstatus)
            if '#noofchildren' in paragraph.text:
                paragraph.text = paragraph.text.replace('#noofchildren', noofchildren)
            if '#childname ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childname', childname)
            if '#childage ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childage', childage)
            if '#childdob' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childdob', childdob)
            if '#childcustody' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childcustody', childcustody)
            if '#section' in paragraph.text:
                paragraph.text = paragraph.text.replace('#section', section)
            if '#act' in paragraph.text:
                paragraph.text = paragraph.text.replace('#act', act)
            if '#caseno' in paragraph.text:
                paragraph.text = paragraph.text.replace('#caseno', caseno)
            if '#idproof' in paragraph.text:
                paragraph.text = paragraph.text.replace('#idproof', idproof)
            if '#marriageproof' in paragraph.text:
                paragraph.text = paragraph.text.replace('#marriageproof', marriageproof)
            if '#residentialproof' in paragraph.text:
                paragraph.text = paragraph.text.replace('#residentialproof', residentialproof)
            if '#propertydocument' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertydocument', propertydocument)
            if '#receipt' in paragraph.text:
                paragraph.text = paragraph.text.replace('#receipt', receipt)
           
            



        return doc
    else:
        return None

@app.route('/print_info_divorce', methods=['POST'])
def print_info_divorce():
    data = request.json
    pet1name = data.get('pet1name')
    pet1age = data.get('pet1age')
    pet1occupation = data.get('pet1occupation')
    pet1address = data.get('pet1address')
    pet1mobileNo = data.get('pet1mobileNo')
    pet1emailid = data.get('pet1emailid')
    pet2name = data.get('pet2name')
    pet2age = data.get('pet2age')
    pet2occupation = data.get('pet2occupation')
    pet2address = data.get('pet2address')
    pet2mobileNo = data.get('pet2mobileNo')
    pet2emailid = data.get('pet2emailid')
    marriedplace = data.get('marriedplace')
    marrieddate = data.get('marrieddate')
    religion = data.get('religion')
    registarplace = data.get('registarplace')
    pet1premarstatus = data.get('pet1premarstatus')
    pet2premarstatus = data.get('pet2premarstatus')
    noofchildren = data.get('noofchildren')
    childname = data.get('childname')
    childage = data.get('childage')
    childdob = data.get('childdob')
    childcustody = data.get('childcustody')
    section = data.get('section')
    act = data.get('act')
    caseno = data.get('caseno')
    idproof = data.get('idproof')
    marriageproof = data.get('marriageproof')
    residentialproof = data.get('residentialproof')
    propertydocument = data.get('propertydocument')
    receipt = data.get('receipt')
    
    
    
    # Update the document with the received values
    updated_doc_Divorce = update_document_Divorce(pet1name,pet1age,pet1occupation,pet1address,pet1mobileNo,pet1emailid,pet2name,pet2age,
                    pet2occupation,pet2address,pet2mobileNo,pet2emailid,marriedplace,marrieddate,religion,registarplace,
                    pet1premarstatus,pet2premarstatus,noofchildren,childname,childage,childdob,childcustody,
                    section,act,caseno,idproof,marriageproof,residentialproof,propertydocument,receipt)
    
    if updated_doc_Divorce:
        # Send the updated document back to the React app
        updated_doc_Divorce.save('updated_document_Divorce.docx')
        return send_file('updated_document_Divorce.docx', as_attachment=True)
    else:
        return "Failed to fetch the document from the URL", 500


def update_document_Sale(date,month,year,vendorname,vendorfathername,vendorreligion,vendorage,vendoraddress,vendoraadharnumber,
                    vendorpannumber,purchasername,purchaserfathername,purchaserreligion,purchaserage,purchaseraddress,purchaseraadharnumber,purchaserpannumber,
                    propertydetails,ownerloanaccountnumber,loansumnumbers,loansumwords,propertytobesold,propertyoldnumber,propertynewnumber,
                    propertypattanumber,propertyaddress,propertyarea,otherpropertiesforsale,saleamountnumbers,saleamountwords,sumadvance1,sum1mode1,sum1mode1date,sum1mode1number,
                    sum1mode1bankname,sum1mode1bankbranch,sumadvance2,sum2mode2,sum2mode2number,sum2mode2bankname,sum2mode2bankbranch,vendorbankname,vendorbankbranch,vendorbanknumber,
                    sumadvance3,sum3mode3,sum3mode3number,sum3mode3bankname,sum3mode3bankbranch,furtherpaymentdetails,directionalmeasurementsinfeet,witness1name,witness2name  ):


    # Download the document from the URL
    document_url = "https://docs.google.com/document/d/1Q2yvGW9i0C0nW-o7UGKO6oquOkP-vXgu/export?format=docx"
    response = requests.get(document_url)
    
    # Check if the request was successful
    if response.status_code == 200:
        # Load the downloaded document
        doc = Document(io.BytesIO(response.content))

        # Replace placeholders with values
        for paragraph in doc.paragraphs:
            if '#date' in paragraph.text:
                paragraph.text = paragraph.text.replace('#date', date)
            if '#month' in paragraph.text:
                paragraph.text = paragraph.text.replace('#month', month)
            if '#year' in paragraph.text:
                paragraph.text = paragraph.text.replace('#year', year)
            if '#vendorname' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendorname', vendorname)
            if '#vendorfathername' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendorfathername', vendorfathername)
            if '#vendorreligion' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendorreligion', vendorreligion)
            if '#vendorage' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendorage', vendorage)
            if '#vendoraddress' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendoraddress', vendoraddress)
            if '#vendoraadharnumber ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendoraadharnumber', vendoraadharnumber)
            if '#vendorpannumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#vendorpannumber', vendorpannumber)
            if '#purchasername' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchasername', purchasername)
            if '#purchaserfathername' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchaserfathername', purchaserfathername)
            if '#purchaserreligion' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchaserreligion', purchaserreligion)
            if '#purchaserage' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchaserage', purchaserage)
            if '#purchaseraddress ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchaseraddress ', purchaseraddress)
            if '#purchaseraadharnumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchaseraadharnumber', purchaseraadharnumber)
            if '#purchaserpannumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#purchaserpannumber', purchaserpannumber)
            if '#propertydetails' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertydetails', propertydetails)
            if '#ownerloanaccountnumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#ownerloanaccountnumber', ownerloanaccountnumber)
            if '#loansumnumbers ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#loansumnumbers', loansumnumbers)
            if '#loansumwords ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#loansumwords', loansumwords)
            if '#propertytobesold' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertytobesold', propertytobesold)
            if '#propertyoldnumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertyoldnumber', propertyoldnumber)
            if '#propertynewnumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertynewnumber', propertynewnumber)
            if '#propertypattanumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertypattanumber', propertypattanumber)
            if '#propertyaddress' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertyaddress', propertyaddress)
            if '#propertyarea' in paragraph.text:
                paragraph.text = paragraph.text.replace('#propertyarea', propertyarea)
            if '#otherpropertiesforsale' in paragraph.text:
                paragraph.text = paragraph.text.replace('#otherpropertiesforsale', otherpropertiesforsale)
            if '#saleamountnumbers' in paragraph.text:
                paragraph.text = paragraph.text.replace('#saleamountnumbers', saleamountnumbers)
            if '#saleamountwords' in paragraph.text:
                paragraph.text = paragraph.text.replace('#saleamountwords', saleamountwords)
            if '#sumadvance1' in paragraph.text:
                paragraph.text = paragraph.text.replace('#sumadvance1', sumadvance1)
            if '#sum1mode1' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum1mode1', sum1mode1)
            if '#sum1mode1date' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum1mode1date', sum1mode1date)
            if '#sum1mode1number' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum1mode1number', sum1mode1number)
            if '#sum1mode1bankname' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum1mode1bankname', sum1mode1bankname)
            if '#sum1mode1bankbranch' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum1mode1bankbranch', sum1mode1bankbranch)
            if '#sumadvance2' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sumadvance2', sumadvance2)
            if '#sum2mode2' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum2mode2', sum2mode2)
            if '#sum2mode2number' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum2mode2number', sum2mode2number)
            if '#sum2mode2bankname' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum2mode2bankname', sum2mode2bankname)
            if '#sum2mode2bankbranch' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum2mode2bankbranch', sum2mode2bankbranch)
            if '#vendorbankname' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#vendorbankname', vendorbankname)
            if '#vendorbankbranch' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#vendorbankbranch', vendorbankbranch)
            if '#vendorbanknumber' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#vendorbanknumber', vendorbanknumber)
            if '#sumadvance3' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sumadvance3', sumadvance3)
            if '#sum3mode3' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum3mode3', sum3mode3)
            if '#sum3mode3number' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum3mode3number', sum3mode3number)
            if '#sum3mode3bankname' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum3mode3bankname', sum3mode3bankname)
            if '#sum3mode3bankbranch' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#sum3mode3bankbranch', sum3mode3bankbranch)
            if '#furtherpaymentdetails' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#furtherpaymentdetails', furtherpaymentdetails)
            if '#directionalmeasurementsinfeet' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#directionalmeasurementsinfeet', directionalmeasurementsinfeet)
            if '#witness1name' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#witness1name', witness1name)
            if '#witness2name' in paragraph.text:
                    paragraph.text = paragraph.text.replace('#witness2name', witness2name)
            
            



        return doc
    else:
        return None

@app.route('/print_info_sale', methods=['POST'])
def print_info_sale():
    data = request.json
    date = data.get('date')
    month = data.get('month')
    year = data.get('year')
    vendorname = data.get('vendorname')
    vendorfathername = data.get('vendorfathername')
    vendorreligion = data.get('vendorreligion')
    vendorage = data.get('vendorage')
    vendoraddress = data.get('vendoraddress')
    vendoraadharnumber = data.get('vendoraadharnumber')
    vendorpannumber = data.get('vendorpannumber')
    purchasername = data.get('purchasername')
    purchaserfathername = data.get('purchaserfathername')
    purchaserreligion = data.get('purchaserreligion')
    purchaserage = data.get('purchaserage')
    purchaseraddress = data.get('purchaseraddress')
    purchaseraadharnumber = data.get('purchaseraadharnumber')
    purchaserpannumber = data.get('purchaserpannumber')
    propertydetails = data.get('propertydetails')
    ownerloanaccountnumber = data.get('ownerloanaccountnumber')
    loansumnumbers = data.get('loansumnumbers')
    loansumwords = data.get('loansumwords')
    propertytobesold = data.get('propertytobesold')
    propertyoldnumber = data.get('propertyoldnumber')
    propertynewnumber = data.get('propertynewnumber')
    propertypattanumber = data.get('propertypattanumber')
    propertyaddress = data.get('propertyaddress')
    propertyarea = data.get('propertyarea')
    otherpropertiesforsale = data.get('otherpropertiesforsale')
    saleamountnumbers = data.get('saleamountnumbers')
    saleamountwords = data.get('saleamountwords')
    sumadvance1 = data.get('sumadvance1')
    sum1mode1 = data.get('sum1mode1')
    sum1mode1date = data.get('sum1mode1date')
    sum1mode1number = data.get('sum1mode1number')
    sum1mode1bankname = data.get('sum1mode1bankname')
    sum1mode1bankbranch = data.get('sum1mode1bankbranch')
    sumadvance2 = data.get('sumadvance2')
    sum2mode2 = data.get('sum2mode2')
    sum2mode2number = data.get('sum2mode2number')
    sum2mode2bankname = data.get('sum2mode2bankname')
    sum2mode2bankbranch = data.get('sum2mode2bankbranch')
    vendorbankname = data.get('vendorbankname')
    vendorbankbranch = data.get('vendorbankbranch')
    vendorbanknumber = data.get('vendorbanknumber')
    sumadvance3 = data.get('sumadvance3')
    sum3mode3 = data.get('sum3mode3')
    sum3mode3number = data.get('sum3mode3number')
    sum3mode3bankname = data.get('sum3mode3bankname')
    sum3mode3bankbranch = data.get('sum3mode3bankbranch')
    furtherpaymentdetails = data.get('furtherpaymentdetails')
    directionalmeasurementsinfeet = data.get('directionalmeasurementsinfeet')
    witness1name = data.get('witness1name')
    witness2name = data.get('witness2name')

    
    
    
    # Update the document with the received values
    updated_doc_Sale = update_document_Sale(date,month,year,vendorname,vendorfathername,vendorreligion,vendorage,vendoraddress,vendoraadharnumber,
                    vendorpannumber,purchasername,purchaserfathername,purchaserreligion,purchaserage,purchaseraddress,purchaseraadharnumber,purchaserpannumber,
                    propertydetails,ownerloanaccountnumber,loansumnumbers,loansumwords,propertytobesold,propertyoldnumber,propertynewnumber,
                    propertypattanumber,propertyaddress,propertyarea,otherpropertiesforsale,saleamountnumbers,saleamountwords,sumadvance1,sum1mode1,sum1mode1date,sum1mode1number,
                    sum1mode1bankname,sum1mode1bankbranch,sumadvance2,sum2mode2,sum2mode2number,sum2mode2bankname,sum2mode2bankbranch,vendorbankname,vendorbankbranch,vendorbanknumber,
                    sumadvance3,sum3mode3,sum3mode3number,sum3mode3bankname,sum3mode3bankbranch,furtherpaymentdetails,directionalmeasurementsinfeet,witness1name,witness2name  )

    if updated_doc_Sale:
        # Send the updated document back to the React app
        updated_doc_Sale.save('updated_document_Sale.docx')
        return send_file('updated_document_Sale.docx', as_attachment=True)
    else:
        return "Failed to fetch the document from the URL", 500


def update_document_Adoption(casenumber,caseyear,childname,childdob,childgender,childreligion,petitionerfathername,petitionerparentnameoffather,petitionermothername,
                    petitioneraddress,respondentfathername,respondentparentnameoffather,petitioner1religionpetitioner1age,petitioner2religion,petitioner2age,advocatename,advocateenrollment,
                    advocateoffice,advocatenumber,respondent1religion,respondent1age,respondentaddress,respondentmothername,respondent2age,
                   childregistrationnumber,registrationdate,handoverdate,adoptiondate,permissiondate,datetoday):


    # Download the document from the URL
    document_url = "https://docs.google.com/document/d/1Ks9idQYdFmV7av99ApTwLsciLJde0GG2/export?format=docx"
    response = requests.get(document_url)
    
    # Check if the request was successful
    if response.status_code == 200:
        # Load the downloaded document
        doc = Document(io.BytesIO(response.content))

        # Replace placeholders with values
        for paragraph in doc.paragraphs:
            if '#casenumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#casenumber', casenumber)
            if '#caseyear' in paragraph.text:
                paragraph.text = paragraph.text.replace('#caseyear', caseyear)
            if '#childname' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childname', childname)
            if '#childdob' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childdob', childdob)
            if '#childgender' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childgender', childgender)
            if '#childreligion' in paragraph.text:
                paragraph.text = paragraph.text.replace('#childreligion', childreligion)
            if '#petitionerfathername' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitionerfathername', petitionerfathername)
            if '#petitionerparentnameoffather' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitionerparentnameoffather', petitionerparentnameoffather)
            if '#petitionermothername  ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitionermothername ', petitionermothername )
            if '#petitioneraddress' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitioneraddress', petitioneraddress)
            if '#respondentfathername' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondentfathername', respondentfathername)
            if '#respondentparentnameoffather' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondentparentnameoffather', respondentparentnameoffather)
            if '#petitioner1religionpetitioner1age ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitioner1religionpetitioner1age ', petitioner1religionpetitioner1age )
            if '#petitioner2religion' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitioner2religion', petitioner2religion)
            if '#petitionerage ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#petitionerage ', petitioner2age)
            if '#advocatename' in paragraph.text:
                paragraph.text = paragraph.text.replace('#advocatename', advocatename)
            if '#advocateenrollment' in paragraph.text:
                paragraph.text = paragraph.text.replace('#advocateenrollment', advocateenrollment)
            if '#advocateoffice' in paragraph.text:
                paragraph.text = paragraph.text.replace('#advocateoffice', advocateoffice)
            if '#advocatenumber' in paragraph.text:
                paragraph.text = paragraph.text.replace('#advocatenumber', advocatenumber)
            if '#respondent1religion ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondent1religion', respondent1religion)
            if '#respondent1age ' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondent1age', respondent1age)
            if '#respondentaddress' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondentaddress', respondentaddress)
            if '#respondentmothername' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondentmothername', respondentmothername)
            if '#respondent2age' in paragraph.text:
                paragraph.text = paragraph.text.replace('#respondent2age', respondent2age)
   
            # if '#childregistrationnumber' in paragraph.text:
            #     paragraph.text = paragraph.text.replace('#childregistrationnumber', childregistrationnumber)
            if '#registrationdate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#registrationdate', registrationdate)
            if '#handoverdate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#handoverdate', handoverdate)
            if '#adoptiondate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#adoptiondate', adoptiondate)
            if '#permissiondate' in paragraph.text:
                paragraph.text = paragraph.text.replace('#permissiondate', permissiondate)
            if '#datetoday' in paragraph.text:
                paragraph.text = paragraph.text.replace('#datetoday', datetoday)
            
            
    
        return doc
    else:
        return None

@app.route('/print_info_real_adoption', methods=['POST'])
def print_info_adoption():
    data = request.json
    casenumber = data.get('casenumber')
    caseyear = data.get('caseyear')
    childname = data.get('childname')
    childdob = data.get('childdob')
    childgender = data.get('childgender')
    childreligion = data.get('childreligion')
    petitionerfathername = data.get('petitionerfathername')
    petitionerparentnameoffather = data.get('petitionerparentnameoffather')
    petitionermothername  = data.get('petitionermothername ')
    petitioneraddress = data.get('petitioneraddress')
    respondentfathername = data.get('respondentfathername')
    respondentparentnameoffather = data.get('respondentparentnameoffather')
    petitioner1religionpetitioner1age  = data.get('petitioner1religionpetitioner1age ')
    petitioner2religion = data.get('petitioner2religion')
    petitioner2age = data.get('petitioner2age')
    advocatename = data.get('advocatename')
    advocateenrollment = data.get('advocateenrollment')
    advocateoffice = data.get('advocateoffice')
    advocatenumber = data.get('advocatenumber')
    respondent1religion = data.get('respondent1religion')
    respondent1age = data.get('respondent1age')
    respondentaddress = data.get('respondentaddress')
    respondentmothername = data.get('respondentmothername')
    respondent2age = data.get('respondent2age')
    # respondentmarriagedate = data.get('childregistrationnumber')
    childregistrationnumber = data.get('propertyaddress')
    registrationdate = data.get('registrationdate')
    handoverdate = data.get('handoverdate')
    adoptiondate = data.get('adoptiondate')
    permissiondate = data.get('permissiondate')
    datetoday = data.get('datetoday')
    

    
    
    
    # Update the document with the received values
    updated_doc_Adoption = update_document_Adoption(casenumber,caseyear,childname,childdob,childgender,childreligion,petitionerfathername,petitionerparentnameoffather,petitionermothername,
                    petitioneraddress,respondentfathername,respondentparentnameoffather,petitioner1religionpetitioner1age,petitioner2religion,petitioner2age,advocatename,advocateenrollment,
                    advocateoffice,advocatenumber,respondent1religion,respondent1age,respondentaddress,respondentmothername,respondent2age,
                   childregistrationnumber,registrationdate,handoverdate,adoptiondate,permissiondate,datetoday)

    if updated_doc_Adoption:
        # Send the updated document back to the React app
        updated_doc_Adoption.save('updated_document_Adoption.docx')
        return send_file('updated_document_Adoption.docx', as_attachment=True)
    else:
        return "Failed to fetch the document from the URL", 500

if __name__ == '__main__':
    app.run(host='0.0.0.0',debug=True, port=5000)
